"""
app.py
------
VicSherlock — Conversation-to-Knowledge AI agent for Vic.ai.
Upload a video recording → get a structured step-by-step guide with screenshots.
Also handles Slack bot events for proactive documentation detection.

Run with: python app.py
Then open http://localhost:5000
"""

import os
import sys
import shutil
import threading
import uuid
import logging
import time
from datetime import datetime
from pathlib import Path

from flask import Flask, render_template, request, jsonify, send_file
from dotenv import load_dotenv
from werkzeug.utils import secure_filename

load_dotenv()

from video_processor import extract_audio, detect_scene_changes, extract_frames_at_timestamps, get_video_duration
from frame_analyzer import process_all_frames, filter_duplicate_frames
from guide_generator import GuideGenerator
from doc_builder import generate_guide_document

from transcriber import transcribe_audio, format_transcript_with_timestamps

try:
    from slack_bot import create_slack_bot_scanner
except ImportError:
    create_slack_bot_scanner = None

# Slack Bolt for interactive messaging
try:
    from slack_bolt import App as SlackBoltApp
    from slack_bolt.adapter.flask import SlackRequestHandler
    from slack_sdk import WebClient
    from anthropic import Anthropic as AnthropicClient
    BOLT_AVAILABLE = True
except ImportError:
    BOLT_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# --- Slack Bolt Setup (for receiving DMs / interactive messages) ---
slack_bolt_app = None
slack_handler = None

if BOLT_AVAILABLE:
    slack_token = os.getenv("SLACK_BOT_TOKEN")
    slack_signing_secret = os.getenv("SLACK_SIGNING_SECRET")
    anthropic_key = os.getenv("ANTHROPIC_API_KEY")

    if slack_token and slack_signing_secret:
        try:
            slack_bolt_app = SlackBoltApp(
                token=slack_token,
                signing_secret=slack_signing_secret,
            )
            slack_handler = SlackRequestHandler(slack_bolt_app)
            logger.info("Slack Bolt app initialized — ready to receive messages")

            # Handle incoming DMs
            @slack_bolt_app.event("message")
            def handle_message_events(body, say, client, event):
                """Respond to DMs sent to VicSherlock."""
                # Ignore bot's own messages
                if event.get("bot_id") or event.get("subtype"):
                    return

                user_id = event.get("user", "")
                user_text = event.get("text", "").strip()
                channel = event.get("channel", "")

                if not user_text:
                    return

                logger.info(f"Received message from {user_id}: {user_text[:100]}")

                try:
                    # Use Claude to generate a helpful response
                    if anthropic_key:
                        anthropic_client = AnthropicClient(api_key=anthropic_key)
                        response = anthropic_client.messages.create(
                            model="claude-sonnet-4-20250514",
                            max_tokens=1024,
                            system="""You are VicSherlock, a Conversation-to-Knowledge AI bot for Vic.ai.
Your job is to monitor Slack conversations for documentation-worthy content and help keep team knowledge up to date.

When users ask you to update documentation or Guru cards, acknowledge the request and explain what you'll do.
When users ask how you work, explain that you scan Slack channels for tutorials, process changes, troubleshooting threads, and other documentation-worthy conversations, then alert the right people to update docs.

Be friendly, concise, and helpful. Use emoji sparingly.""",
                            messages=[{"role": "user", "content": user_text}],
                        )
                        reply = response.content[0].text
                    else:
                        reply = "Hey! I'm VicSherlock. I scan Slack for documentation-worthy conversations. I can't process your request right now, but I'm here to help!"

                    say(text=reply)
                    logger.info(f"Replied to {user_id} in {channel}")

                except Exception as e:
                    logger.error(f"Error responding to message: {e}")
                    say(text="Oops, I hit a snag processing that. Try again in a moment!")

            # Handle app_home_opened event (optional)
            @slack_bolt_app.event("app_home_opened")
            def handle_app_home(event, client):
                """Update the App Home tab when a user opens it."""
                try:
                    client.views_publish(
                        user_id=event["user"],
                        view={
                            "type": "home",
                            "blocks": [
                                {
                                    "type": "header",
                                    "text": {"type": "plain_text", "text": "VicSherlock"}
                                },
                                {
                                    "type": "section",
                                    "text": {
                                        "type": "mrkdwn",
                                        "text": "*Your Conversation-to-Knowledge AI Agent*\n\nI scan Slack channels for documentation-worthy conversations — tutorials, process changes, troubleshooting threads — and alert you when it's time to update your docs."
                                    }
                                },
                                {"type": "divider"},
                                {
                                    "type": "section",
                                    "text": {
                                        "type": "mrkdwn",
                                        "text": "*What I can do:*\n• Scan channels for knowledge worth capturing\n• Alert you when Guru cards need updating\n• Convert video recordings into step-by-step guides\n• Help keep your team's documentation current"
                                    }
                                },
                                {
                                    "type": "section",
                                    "text": {
                                        "type": "mrkdwn",
                                        "text": "Send me a message to get started!"
                                    }
                                },
                            ]
                        }
                    )
                except Exception as e:
                    logger.error(f"Error publishing home tab: {e}")

        except Exception as e:
            logger.error(f"Failed to initialize Slack Bolt: {e}")
            slack_bolt_app = None
            slack_handler = None
app.config["MAX_CONTENT_LENGTH"] = 500 * 1024 * 1024  # 500MB max upload

# Directories
UPLOAD_DIR = Path("./uploads")
TEMP_DIR = Path("./temp")
OUTPUT_DIR = Path("./output")

for d in [UPLOAD_DIR, TEMP_DIR, OUTPUT_DIR]:
    d.mkdir(exist_ok=True)

# Job tracking
jobs = {}

ALLOWED_EXTENSIONS = {".mp4", ".mov", ".avi", ".mkv", ".webm"}


def allowed_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def run_video_job(job_id: str, video_path: str, instructions: str, doc_type: str, transcript: str = None):
    """Full pipeline: video → audio → transcript → frames → guide → docx.

    If transcript is provided, skip audio extraction and transcription.
    If transcript is None, extract audio and transcribe locally using faster-whisper.
    """
    job = jobs[job_id]
    job_temp = TEMP_DIR / job_id
    job_temp.mkdir(parents=True, exist_ok=True)

    try:
        # --- Step 1 & 2: Handle transcription ---
        if transcript:
            # Use provided transcript
            job["step"] = "Using provided transcript..."
            job["progress"] = 25
            logger.info(f"Job {job_id}: Using provided transcript")
            transcript_text = transcript
            segments = []
        else:
            # Extract audio and transcribe locally
            job["step"] = "Extracting audio from video..."
            job["progress"] = 10
            logger.info(f"Job {job_id}: Extracting audio")

            audio_path = str(job_temp / "audio.wav")
            extract_audio(video_path, audio_path)

            # Transcribe with local Whisper model
            job["step"] = "Transcribing audio with Whisper..."
            job["progress"] = 25
            logger.info(f"Job {job_id}: Transcribing")

            transcript_result = transcribe_audio(audio_path)
            transcript_text = transcript_result["text"]
            segments = transcript_result.get("segments", [])

        job["transcript_preview"] = transcript_text[:500] + "..." if len(transcript_text) > 500 else transcript_text

        # --- Step 3: Extract screenshots ---
        job["step"] = "Extracting screenshots from video..."
        job["progress"] = 45
        logger.info(f"Job {job_id}: Extracting frames")

        raw_frames_dir = str(job_temp / "frames_raw")
        timestamps = detect_scene_changes(video_path, max_frames=15)
        raw_frames = extract_frames_at_timestamps(video_path, timestamps, raw_frames_dir)

        # --- Step 4: Crop faces from screenshots ---
        job["step"] = "Processing screenshots (removing faces)..."
        job["progress"] = 55
        logger.info(f"Job {job_id}: Processing frames — cropping faces")

        clean_frames_dir = str(job_temp / "frames_clean")
        cleaned_frames = process_all_frames(raw_frames_dir, clean_frames_dir)

        # Remove near-duplicate frames
        unique_frames = filter_duplicate_frames(cleaned_frames)
        logger.info(f"Job {job_id}: {len(unique_frames)} unique frames after filtering")

        # --- Step 5: Generate guide with Claude ---
        job["step"] = "Generating guide with Claude..."
        job["progress"] = 70
        logger.info(f"Job {job_id}: Generating guide")

        anthropic_key = os.getenv("ANTHROPIC_API_KEY")
        if not anthropic_key:
            raise RuntimeError("Missing ANTHROPIC_API_KEY in environment variables.")

        generator = GuideGenerator(anthropic_key)
        guide = generator.generate_guide(transcript_text, instructions, doc_type)

        # Map screenshots to steps
        frame_map = generator.map_screenshots_to_steps(guide["steps"], unique_frames)

        # --- Step 6: Generate .docx ---
        job["step"] = "Building document with screenshots..."
        job["progress"] = 90
        logger.info(f"Job {job_id}: Building docx")

        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        safe_title = "".join(c for c in guide["title"] if c.isalnum() or c in " _-")[:50].strip()
        filename = f"VicSherlock_{safe_title}_{timestamp}.docx"
        output_path = str(OUTPUT_DIR / filename)

        generate_guide_document(guide, frame_map, output_path)

        # --- Done ---
        job["status"] = "done"
        job["step"] = "Complete!"
        job["progress"] = 100
        job["file"] = {"name": guide["title"], "filename": filename}
        job["guide_title"] = guide["title"]
        job["step_count"] = len(guide["steps"])
        job["screenshot_count"] = len(unique_frames)

        logger.info(f"Job {job_id}: Complete — {filename}")

    except Exception as e:
        logger.error(f"Job {job_id}: Failed — {str(e)}")
        job["status"] = "error"
        job["error"] = str(e)

    finally:
        # Cleanup temp files
        try:
            shutil.rmtree(str(job_temp), ignore_errors=True)
            if os.path.exists(video_path):
                os.remove(video_path)
        except Exception:
            pass


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload():
    # Validate file
    if "video" not in request.files:
        return jsonify({"error": "No video file provided."}), 400

    file = request.files["video"]
    if file.filename == "" or not allowed_file(file.filename):
        return jsonify({"error": "Please upload an MP4, MOV, AVI, MKV, or WebM video file."}), 400

    instructions = request.form.get("instructions", "").strip()
    doc_type = request.form.get("doc_type", "step-by-step").strip()
    transcript = request.form.get("transcript", "").strip()

    if not instructions:
        return jsonify({"error": "Please provide instructions for the guide."}), 400

    # Save uploaded file
    job_id = str(uuid.uuid4())
    filename = secure_filename(file.filename)
    video_path = str(UPLOAD_DIR / f"{job_id}_{filename}")
    file.save(video_path)

    # Create job
    jobs[job_id] = {
        "status": "running",
        "step": "Starting...",
        "progress": 0,
        "file": None,
        "error": None,
    }

    # Run in background
    thread = threading.Thread(
        target=run_video_job,
        args=(job_id, video_path, instructions, doc_type, transcript if transcript else None)
    )
    thread.daemon = True
    thread.start()

    return jsonify({"job_id": job_id})


@app.route("/status/<job_id>")
def status(job_id):
    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found."}), 404
    return jsonify(job)


@app.route("/download/<path:filename>")
def download(filename):
    file_path = OUTPUT_DIR / filename
    if not file_path.exists():
        return "File not found.", 404
    return send_file(str(file_path), as_attachment=True, download_name=filename)


@app.route("/scan-channels", methods=["POST"])
def scan_channels():
    """Scan Slack channels for documentation-worthy conversations."""
    try:
        if not create_slack_bot_scanner:
            return jsonify({"error": "Slack bot module not available."}), 400
        bot = create_slack_bot_scanner()
        if not bot:
            return jsonify({
                "error": "Slack bot not configured. Please set SLACK_BOT_TOKEN and ANTHROPIC_API_KEY."
            }), 400

        # Get optional limit parameter
        limit = request.args.get("limit", type=int)

        # Run scan in background
        job_id = str(uuid.uuid4())
        scan_job = {
            "status": "running",
            "step": "Initializing Slack channel scan...",
            "progress": 0,
            "results": None,
            "error": None,
        }
        jobs[job_id] = scan_job

        def run_slack_scan():
            try:
                scan_job["step"] = "Fetching channels..."
                scan_job["progress"] = 10

                scan_job["step"] = "Analyzing conversations with Claude..."
                scan_job["progress"] = 30

                results = bot.perform_full_scan_and_notify(limit_channels=limit)

                scan_job["step"] = "Complete!"
                scan_job["progress"] = 100
                scan_job["status"] = "done"
                scan_job["results"] = results

                logger.info(f"Slack scan complete: {results}")

            except Exception as e:
                logger.error(f"Error in Slack scan: {str(e)}")
                scan_job["status"] = "error"
                scan_job["error"] = str(e)

        thread = threading.Thread(target=run_slack_scan)
        thread.daemon = True
        thread.start()

        return jsonify({"job_id": job_id, "status": "scanning"})

    except Exception as e:
        logger.error(f"Error starting Slack scan: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route("/scan-status/<job_id>")
def scan_status(job_id):
    """Get status of a Slack scan job."""
    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found."}), 404

    # Return scan-specific fields
    return jsonify({
        "status": job.get("status"),
        "step": job.get("step"),
        "progress": job.get("progress"),
        "results": job.get("results"),
        "error": job.get("error"),
    })


# --- Slack Events Endpoint (for Bolt to receive messages) ---
@app.route("/slack/events", methods=["POST"])
def slack_events():
    """Handle Slack events directly (no Bolt signature check for demo)."""
    data = request.get_json(silent=True)
    if not data:
        return jsonify({"error": "no data"}), 400

    # Handle Slack URL verification challenge
    if data.get("type") == "url_verification":
        logger.info("Received Slack URL verification challenge")
        return jsonify({"challenge": data.get("challenge", "")})

    # Handle event callbacks directly (bypass Bolt for reliability)
    if data.get("type") == "event_callback":
        event = data.get("event", {})
        event_type = event.get("type", "")

        logger.info(f"Received Slack event: {event_type}")

        # Handle DM messages
        if event_type == "message" and not event.get("bot_id") and not event.get("subtype"):
            user_text = event.get("text", "").strip()
            channel = event.get("channel", "")
            user_id = event.get("user", "")
            files = event.get("files", [])

            if (user_text or files) and channel:
                # Process in background thread so we respond to Slack within 3s
                def reply_with_claude(msg_text=user_text, msg_files=files, msg_channel=channel, msg_user=user_id):
                    try:
                        anthropic_key = os.getenv("ANTHROPIC_API_KEY")
                        slack_token = os.getenv("SLACK_BOT_TOKEN")
                        if not anthropic_key or not slack_token:
                            logger.error("Missing ANTHROPIC_API_KEY or SLACK_BOT_TOKEN")
                            return

                        from anthropic import Anthropic as AnthropicClient
                        from slack_sdk import WebClient
                        import requests as req

                        client = AnthropicClient(api_key=anthropic_key)
                        slack = WebClient(token=slack_token)

                        # Check if user attached a PDF file
                        pdf_text = None
                        pdf_filename = None
                        for f in msg_files:
                            if f.get("filetype") == "pdf" or (f.get("name", "").lower().endswith(".pdf")):
                                pdf_filename = f.get("name", "document.pdf")
                                pdf_url = f.get("url_private_download") or f.get("url_private")
                                if pdf_url:
                                    logger.info(f"Downloading PDF attachment: {pdf_filename}")
                                    headers = {"Authorization": f"Bearer {slack_token}"}
                                    resp = req.get(pdf_url, headers=headers)
                                    if resp.status_code == 200:
                                        # Extract text from PDF
                                        try:
                                            import io
                                            from PyPDF2 import PdfReader
                                            reader = PdfReader(io.BytesIO(resp.content))
                                            pages = []
                                            for page in reader.pages:
                                                text = page.extract_text()
                                                if text:
                                                    pages.append(text)
                                            pdf_text = "\n\n".join(pages)
                                            logger.info(f"Extracted {len(pages)} pages from PDF ({len(pdf_text)} chars)")
                                        except Exception as e:
                                            logger.error(f"PDF extraction error: {e}")
                                break  # Only process first PDF

                        vic_system_prompt = """You are VicSherlock, a Conversation-to-Knowledge AI bot for Vic.ai.
Your job is to monitor Slack conversations for documentation-worthy content and help keep team knowledge up to date.

You actively scan Slack channels for tutorials, process changes, troubleshooting threads, and other documentation-worthy conversations, then alert the right people to update docs.
You can also convert video recordings into step-by-step implementation guides with screenshots at https://vicsherlock.onrender.com

RECENT FINDINGS FROM YOUR CHANNEL SCANS:

1. In #project_vicpaygolive, Katie Roy (on Jan 7, 2026) announced a process change:
   - The implementations team will now own uploading the logo and signature for VicPay implementations
   - The ability to upload the signature and logo is available in VicAdmin
   - She mentioned she will update the Guru card, but is offboarding that responsibility to the implementation team
   - The relevant Guru card is: "Implementing Payments (VicPay)"
   - Action needed: Update the Implementing Payments (VicPay) doc in Guru to reflect the new ownership of logo/signature uploads

2. You also found a VicPay implementation training recording where a team member walked through:
   - Setting up the cash account / GL account in Sage Intacct UAT
   - Creating a sub-ledger writer for the cash account
   - Enabling payments in UAT
   - Configuring VicPay settings

When users ask about recent findings or Katie Roy's messages, share these details.
When users ask you to update documentation or Guru cards, acknowledge the request and explain what steps are needed.

Be friendly, concise, and helpful. Use emoji sparingly."""

                        # If a PDF was attached, build a doc-update flow
                        if pdf_text and msg_text:
                            # Tell user we're working on it
                            slack.chat_postMessage(
                                channel=msg_channel,
                                text="Got it! I'm reading the document and updating it with the latest findings. Give me a moment..."
                            )

                            # Ask Claude to produce the updated document content
                            update_response = client.messages.create(
                                model="claude-sonnet-4-20250514",
                                max_tokens=4096,
                                system=vic_system_prompt + """

IMPORTANT: The user has shared a document and wants you to update it.
You have the full text of the document below. Apply the relevant changes based on your recent findings from channel scans.
Return the FULL updated document text with all changes incorporated.
Mark any new or changed sections with [UPDATED] at the start of the line so changes are easy to spot.
Keep the existing structure and formatting. Only modify sections that need updating based on your findings.""",
                                messages=[{"role": "user", "content": f"Here is the document '{pdf_filename}':\n\n{pdf_text}\n\nUser request: {msg_text}"}],
                            )
                            updated_content = update_response.content[0].text

                            # Generate a .docx with the updated content
                            try:
                                from docx import Document
                                from docx.shared import Pt, Inches
                                import tempfile

                                doc = Document()
                                doc.core_properties.author = "VicSherlock"

                                # Parse the updated content into the docx
                                lines = updated_content.split("\n")
                                for line in lines:
                                    stripped = line.strip()
                                    if not stripped:
                                        continue
                                    if stripped.startswith("# "):
                                        doc.add_heading(stripped[2:], level=1)
                                    elif stripped.startswith("## "):
                                        doc.add_heading(stripped[3:], level=2)
                                    elif stripped.startswith("### "):
                                        doc.add_heading(stripped[4:], level=3)
                                    elif stripped.startswith("- ") or stripped.startswith("* "):
                                        doc.add_paragraph(stripped[2:], style="List Bullet")
                                    elif stripped.startswith("[UPDATED]"):
                                        p = doc.add_paragraph()
                                        run = p.add_run(stripped)
                                        run.bold = True
                                    else:
                                        doc.add_paragraph(stripped)

                                # Save to temp file
                                base_name = pdf_filename.rsplit(".", 1)[0] if "." in pdf_filename else pdf_filename
                                output_name = f"{base_name} - Updated by VicSherlock.docx"
                                tmp_path = os.path.join(tempfile.gettempdir(), output_name)
                                doc.save(tmp_path)

                                # Upload to Slack
                                slack.files_upload_v2(
                                    channel=msg_channel,
                                    file=tmp_path,
                                    filename=output_name,
                                    title=output_name,
                                    initial_comment="Here's the updated document with the latest changes incorporated. Sections marked [UPDATED] show where I made changes based on Katie Roy's process update."
                                )
                                logger.info(f"Uploaded updated doc to {msg_channel}")

                                # Clean up
                                os.remove(tmp_path)

                            except Exception as e:
                                logger.error(f"Error generating docx: {e}")
                                # Fall back to sending the text
                                # Truncate if too long for Slack (4000 char limit)
                                if len(updated_content) > 3900:
                                    slack.chat_postMessage(
                                        channel=msg_channel,
                                        text="I updated the document but couldn't generate a .docx file. Here are the key changes:\n\n" + updated_content[:3900] + "\n\n_(truncated — full doc was too long for a message)_"
                                    )
                                else:
                                    slack.chat_postMessage(
                                        channel=msg_channel,
                                        text="Here's the updated document content:\n\n" + updated_content
                                    )
                        else:
                            # Normal text-only conversation
                            user_content = msg_text or "Hello"
                            if pdf_text:
                                user_content = f"[User shared a PDF: {pdf_filename}]\n\nDocument content:\n{pdf_text[:3000]}\n\n{msg_text or 'What can you tell me about this document?'}"

                            response = client.messages.create(
                                model="claude-sonnet-4-20250514",
                                max_tokens=1024,
                                system=vic_system_prompt,
                                messages=[{"role": "user", "content": user_content}],
                            )
                            reply = response.content[0].text
                            slack.chat_postMessage(channel=msg_channel, text=reply)

                        logger.info(f"Replied to {msg_user} in {msg_channel}")

                    except Exception as e:
                        logger.error(f"Error replying to message: {e}")
                        try:
                            from slack_sdk import WebClient
                            slack = WebClient(token=os.getenv("SLACK_BOT_TOKEN"))
                            slack.chat_postMessage(channel=msg_channel, text="Oops, I hit a snag processing that. Try again in a moment!")
                        except Exception:
                            pass

                thread = threading.Thread(target=reply_with_claude)
                thread.daemon = True
                thread.start()

        # Handle app_home_opened
        elif event_type == "app_home_opened":
            user_id = event.get("user", "")
            try:
                from slack_sdk import WebClient
                slack = WebClient(token=os.getenv("SLACK_BOT_TOKEN"))
                slack.views_publish(
                    user_id=user_id,
                    view={
                        "type": "home",
                        "blocks": [
                            {"type": "header", "text": {"type": "plain_text", "text": "VicSherlock"}},
                            {"type": "section", "text": {"type": "mrkdwn", "text": "*Your Conversation-to-Knowledge AI Agent*\n\nI scan Slack channels for documentation-worthy conversations — tutorials, process changes, troubleshooting threads — and alert you when it's time to update your docs."}},
                            {"type": "divider"},
                            {"type": "section", "text": {"type": "mrkdwn", "text": "*What I can do:*\n• Scan channels for knowledge worth capturing\n• Alert you when Guru cards need updating\n• Convert video recordings into step-by-step guides\n• Help keep your team's documentation current"}},
                            {"type": "section", "text": {"type": "mrkdwn", "text": "Send me a message to get started!"}},
                        ]
                    }
                )
            except Exception as e:
                logger.error(f"Error publishing home tab: {e}")

        return jsonify({"ok": True}), 200

    return jsonify({"ok": True}), 200


# --- Background Channel Scanner ---
# Periodically scans Slack channels and posts findings to #vicsherlock

SCAN_INTERVAL_SECONDS = int(os.getenv("SCAN_INTERVAL", "1800"))  # Default: every 30 minutes
_scanner_thread = None
_last_scan_results = None
_scan_lock = threading.Lock()


def run_channel_scan():
    """Execute a single channel scan and post findings to #vicsherlock."""
    global _last_scan_results
    try:
        if create_slack_bot_scanner is None:
            logger.warning("slack_bot module not available — skipping scan")
            return None

        scanner = create_slack_bot_scanner()
        if scanner is None:
            logger.warning("Could not create scanner (missing tokens) — skipping scan")
            return None

        logger.info("Starting scheduled channel scan...")
        results = scanner.perform_full_scan_and_notify(limit_channels=20)

        with _scan_lock:
            _last_scan_results = {
                "timestamp": datetime.now().isoformat(),
                "results": results,
            }

        logger.info(
            f"Scan complete: {results.get('channels_scanned', 0)} channels scanned, "
            f"{len(results.get('documentation_worthy', []))} findings, "
            f"{results.get('notifications', {}).get('sent', 0)} DMs sent"
        )
        return results

    except Exception as e:
        logger.error(f"Channel scan failed: {e}")
        return None


def background_scanner_loop():
    """Background loop that runs channel scans on an interval."""
    # Wait 30 seconds after startup before first scan (let the app warm up)
    time.sleep(30)
    logger.info(f"Background scanner started — scanning every {SCAN_INTERVAL_SECONDS}s")

    while True:
        try:
            run_channel_scan()
        except Exception as e:
            logger.error(f"Background scanner error: {e}")

        time.sleep(SCAN_INTERVAL_SECONDS)


def start_background_scanner():
    """Start the background scanner thread if not already running."""
    global _scanner_thread
    if _scanner_thread is not None and _scanner_thread.is_alive():
        return

    _scanner_thread = threading.Thread(target=background_scanner_loop, daemon=True)
    _scanner_thread.start()
    logger.info("Background channel scanner thread launched")


@app.route("/scan-now", methods=["POST"])
def trigger_scan():
    """Manually trigger a channel scan (POST /scan-now)."""
    def do_scan():
        run_channel_scan()

    thread = threading.Thread(target=do_scan, daemon=True)
    thread.start()
    return jsonify({"status": "scan_started", "message": "Channel scan triggered — findings will be DM'd to you"}), 200


@app.route("/scan-status", methods=["GET"])
def scan_status():
    """Get the last scan results."""
    with _scan_lock:
        if _last_scan_results:
            return jsonify(_last_scan_results), 200
        else:
            return jsonify({"status": "no_scans_yet"}), 200


@app.route("/send-finding", methods=["POST"])
def send_finding():
    """Send a test finding message as VicSherlock bot to the DM channel."""
    slack_token = os.getenv("SLACK_BOT_TOKEN")
    if not slack_token:
        return jsonify({"error": "no SLACK_BOT_TOKEN"}), 500

    data = request.get_json() or {}
    channel_name = data.get("channel_name", "#customer_stonewall-kitchen")
    author = data.get("author", "Anthony Margaritondo")
    topic = data.get("topic", "Approval Flow Behavior When Invoice Fields Are Modified")
    preview = data.get("preview", "Looking into how approvers change if someone modifies an invoice and changes a field that affects the approval flow...")
    target_channel = data.get("target_channel", os.getenv("NOTIFY_CHANNEL_ID", "D09JP0H2DSN"))

    try:
        from slack_sdk import WebClient
        slack = WebClient(token=slack_token)

        message = (
            f"Hey! I just scanned your Slack channels and noticed something in *{channel_name}* from *{author}* that looks like it should be documented:\n\n"
            f"*{topic}*\n\n"
            f"> {preview}\n\n"
            "Would you like me to:\n"
            "• *Create a new doc* from this conversation?\n"
            "• *Update an existing doc* — just send me the current doc/PDF and I'll revise it!\n\n"
            "Just reply here and let me know!"
        )

        result = slack.chat_postMessage(channel=target_channel, text=message, mrkdwn=True)
        return jsonify({"ok": True, "ts": result["ts"]}), 200

    except Exception as e:
        logger.error(f"Error sending finding: {e}")
        return jsonify({"error": str(e)}), 500


# Start the background scanner when the app loads
start_background_scanner()


if __name__ == "__main__":
    print("\n🔍 VicSherlock — Conversation-to-Knowledge AI")
    print("   Open http://localhost:5000 in your browser\n")
    app.run(debug=False, port=5000, host="0.0.0.0")
