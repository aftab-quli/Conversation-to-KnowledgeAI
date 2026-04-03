"""
VicSherlock — Slack Bot
Vic.ai AI Hackathon 2025 · Built by Aftab Quli

This bot:
  - Monitors every Slack channel and DMs the author when it finds something doc-worthy
  - Accepts on-demand DMs: "Create a FAQ from this transcript" / "Update [Guru link] based on this"
  - Handles Zoom recording.completed webhooks — auto-analyzes every cloud-recorded call
  - Generates Vic-branded .docx files and sends them back in the DM thread for review
"""

import os, io, json, re, requests
from flask import Flask, request, jsonify
from slack_bolt import App
from slack_bolt.adapter.flask import SlackRequestHandler
from anthropic import Anthropic

# ── App setup ─────────────────────────────────────────────────────────────────

flask_app = Flask(__name__)

slack_app = App(
    token=os.environ.get("SLACK_BOT_TOKEN"),
    signing_secret=os.environ.get("SLACK_SIGNING_SECRET"),
)
handler = SlackRequestHandler(slack_app)
claude  = Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))

# ── In-memory conversation state (per user DM thread) ────────────────────────
# state[user_id] = {
#   "step":      "awaiting_action" | "awaiting_doc_type" | "awaiting_approval",
#   "content":   <source text>,
#   "source":    "slack" | "zoom" | "gong",
#   "action":    "new" | "update",
#   "doc_type":  "FAQ" | "step-by-step guide" | ...,
#   "update_url": <existing article URL if updating>,
#   "doc_data":  <generated doc JSON>,
# }
state = {}

DOC_TYPES = ["step-by-step guide", "faq", "release notes",
             "troubleshooting guide", "process doc", "training doc", "customer-facing guide"]

# ── Prompts ───────────────────────────────────────────────────────────────────

ANALYZE_PROMPT = """You are VicSherlock, Vic.ai's documentation agent.
Analyze this Slack message and determine if it describes something worth documenting:
a process change, product update, feature flag change, workflow update, or important announcement.

Respond ONLY with a valid JSON object — no markdown, no explanation:
{
  "worthy": true or false,
  "reason": "one sentence explanation",
  "suggested_type": "step-by-step guide | FAQ | release notes | troubleshooting guide | process doc",
  "summary": "one sentence summary of what this is about"
}"""

DOC_PROMPT = """You are VicSherlock, Vic.ai's documentation agent.
Generate a {doc_type} based on the content provided.

Output ONLY a valid JSON object — no markdown, no explanation:
{{
  "title": "Clear, descriptive document title",
  "summary": "1-2 sentence summary",
  "sections": [
    {{
      "heading": "Section heading",
      "content": "Optional intro sentence",
      "steps": ["Clear step or point", "Another step"]
    }}
  ],
  "notes": ["Important caveat or note"]
}}"""

# ── Helpers ───────────────────────────────────────────────────────────────────

def call_claude(system, user_msg, max_tokens=4096):
    resp = claude.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=max_tokens,
        system=system,
        messages=[{"role": "user", "content": user_msg}],
    )
    return resp.content[0].text


def is_doc_worthy(message_text):
    """Ask Claude if a Slack message is worth documenting. Returns dict."""
    try:
        raw  = call_claude(ANALYZE_PROMPT, message_text, max_tokens=512)
        s, e = raw.find("{"), raw.rfind("}") + 1
        return json.loads(raw[s:e])
    except Exception:
        return {"worthy": False}


def generate_doc_json(content, doc_type, source="slack"):
    """Ask Claude to generate a structured doc. Returns dict."""
    system = DOC_PROMPT.format(doc_type=doc_type)
    user   = f"Source: {source}\n\nContent:\n{content}"
    raw    = call_claude(system, user)
    s, e   = raw.find("{"), raw.rfind("}") + 1
    return json.loads(raw[s:e])


def build_docx_bytes(doc_data):
    """Build a Vic-branded .docx from doc_data dict. Returns bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    NAVY   = (28, 32, 67)
    ACCENT = (91, 95, 207)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1)
        sec.left_margin = sec.right_margin = Inches(1.2)

    def run(para, text, bold=False, italic=False, size=11, color=None):
        r = para.add_run(text)
        r.bold, r.italic = bold, italic
        r.font.size = Pt(size)
        r.font.name = "Calibri"
        if color:
            r.font.color.rgb = RGBColor(*color)
        return r

    run(doc.add_paragraph(), doc_data.get("title", "VicSherlock Guide"),
        bold=True, size=22, color=NAVY)

    if doc_data.get("summary"):
        run(doc.add_paragraph(), doc_data["summary"], italic=True, size=11, color=(100,110,145))

    doc.add_paragraph()

    for section in doc_data.get("sections", []):
        run(doc.add_paragraph(), section.get("heading",""), bold=True, size=14, color=ACCENT)
        if section.get("content"):
            run(doc.add_paragraph(), section["content"], size=11)
        for step in section.get("steps", []):
            run(doc.add_paragraph(style="List Number"), step, size=11)
        doc.add_paragraph()

    if doc_data.get("notes"):
        run(doc.add_paragraph(), "Notes", bold=True, size=13, color=NAVY)
        for note in doc_data["notes"]:
            run(doc.add_paragraph(style="List Bullet"), note, size=10)

    doc.add_paragraph()
    run(doc.add_paragraph(),
        "Generated by VicSherlock · Vic.ai · Conversation-to-Knowledge AI",
        italic=True, size=9, color=(150,155,185))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def send_doc_to_user(user_id, doc_data, say_fn, thread_ts=None):
    """Generate .docx and upload it to the user's DM."""
    doc_bytes = build_docx_bytes(doc_data)
    title     = doc_data.get("title", "VicSherlock_Output")
    filename  = title.replace(" ", "_")[:60] + ".docx"

    slack_app.client.files_upload_v2(
        channels=user_id,
        filename=filename,
        content=doc_bytes,
        title=title,
        initial_comment=(
            f"Here's your doc: *{title}*\n\n"
            "Review it and let me know:\n"
            "• Reply *approved* and I'll publish it\n"
            "• Reply *update: [your feedback]* and I'll revise it\n"
            "• Or paste the link to the Guru/Intercom article and I'll update it directly"
        ),
    )


def parse_doc_type(text):
    """Extract doc type from free-text user message."""
    text_lower = text.lower()
    for dt in DOC_TYPES:
        if dt in text_lower:
            return dt
    if "faq"            in text_lower: return "FAQ"
    if "step"           in text_lower: return "step-by-step guide"
    if "release"        in text_lower: return "release notes"
    if "troubleshoot"   in text_lower: return "troubleshooting guide"
    if "training"       in text_lower: return "training doc"
    if "customer"       in text_lower: return "customer-facing guide"
    return None


def extract_url(text):
    """Pull the first URL out of a message."""
    match = re.search(r'https?://\S+', text)
    return match.group(0) if match else None


# ── Channel monitoring ────────────────────────────────────────────────────────

@slack_app.event("message")
def handle_channel_message(event, say, client):
    """Listen to every channel message. If doc-worthy, DM the author."""

    # Ignore bot messages, DMs, and edited messages
    if event.get("bot_id"):                    return
    if event.get("channel_type") == "im":      return
    if event.get("subtype") == "message_changed": return

    text    = event.get("text", "").strip()
    user_id = event.get("user")

    if not text or not user_id or len(text) < 40:
        return

    # Check if doc-worthy
    result = is_doc_worthy(text)
    if not result.get("worthy"):
        return

    # Get channel name for context
    try:
        ch_info  = client.conversations_info(channel=event["channel"])
        ch_name  = ch_info["channel"]["name"]
    except Exception:
        ch_name = "a Slack channel"

    # Initialise state for this user
    state[user_id] = {
        "step":    "awaiting_action",
        "content": text,
        "source":  "slack",
        "channel": ch_name,
        "suggested_type": result.get("suggested_type", "step-by-step guide"),
    }

    # DM the author
    client.chat_postMessage(
        channel=user_id,
        text=(
            f"Hey — VicSherlock here. I was scanning #{ch_name} and spotted something "
            f"that looks worth documenting.\n\n"
            f"*What I found:* {result.get('summary', text[:200])}\n\n"
            f"What would you like to do?\n"
            f"• Reply *new* to create a new guide\n"
            f"• Reply *update* + paste the existing article link to update it\n"
            f"• Reply *skip* to ignore this one\n\n"
            f"I'd suggest a *{result.get('suggested_type', 'guide')}* — "
            f"but you can ask for any type: FAQ, release notes, troubleshooting guide, etc."
        ),
    )


# ── DM handling ───────────────────────────────────────────────────────────────

@slack_app.event("app_mention")
def handle_mention(event, say):
    say("Hey! DM me directly to get started. You can say things like:\n"
        "• *Create a FAQ from this transcript* (then paste it)\n"
        "• *Update [Guru link] based on today's call*\n"
        "• *Write a step-by-step guide from the Zoom call on April 2*")


@slack_app.message("")
def handle_dm(message, say, client):
    """Handle all DMs to the bot."""

    if message.get("channel_type") != "im": return
    if message.get("bot_id"):               return

    user_id = message["user"]
    text    = message.get("text", "").strip()
    text_lo = text.lower()
    s       = state.get(user_id, {})

    # ── SKIP ──
    if text_lo in ("skip", "no", "ignore", "not now"):
        state.pop(user_id, None)
        say("Got it — skipping this one. I'll keep watching.")
        return

    # ── APPROVAL ──
    if s.get("step") == "awaiting_approval" and text_lo.startswith("approved"):
        url = extract_url(text) or s.get("update_url")
        if url:
            say(f"Publishing to {url}... *(auto-publish via API coming soon — "
                f"for now, use the downloaded doc to update the article manually)*")
        else:
            say("Approved! The doc is ready to publish. "
                "Paste the Guru or Intercom article link and I'll push it directly *(coming soon)*.")
        state.pop(user_id, None)
        return

    # ── REVISION REQUEST ──
    if s.get("step") == "awaiting_approval" and text_lo.startswith("update:"):
        feedback    = text[7:].strip()
        revised_doc = generate_doc_json(
            s["content"] + f"\n\nRevision requested: {feedback}",
            s.get("doc_type", "guide"), s.get("source", "slack")
        )
        state[user_id]["doc_data"] = revised_doc
        say("Revised — here's the updated version:")
        send_doc_to_user(user_id, revised_doc, say)
        return

    # ── AWAITING ACTION (bot asked new/update/skip) ──
    if s.get("step") == "awaiting_action":
        if "update" in text_lo:
            url = extract_url(text)
            state[user_id].update({"step": "awaiting_doc_type", "action": "update", "update_url": url})
            say(f"Got it — I'll update the existing article{' at ' + url if url else ''}.\n"
                f"What type of doc is it? (step-by-step guide, FAQ, release notes, troubleshooting guide, etc.) "
                f"Or just say *go* and I'll use my suggestion: *{s.get('suggested_type', 'guide')}*")
        elif "new" in text_lo:
            state[user_id].update({"step": "awaiting_doc_type", "action": "new"})
            say(f"Creating a new guide. What type? (step-by-step guide, FAQ, release notes, troubleshooting guide, etc.) "
                f"Or just say *go* to use my suggestion: *{s.get('suggested_type', 'guide')}*")
        return

    # ── AWAITING DOC TYPE ──
    if s.get("step") == "awaiting_doc_type":
        doc_type = parse_doc_type(text) or s.get("suggested_type", "step-by-step guide")
        if text_lo == "go":
            doc_type = s.get("suggested_type", "step-by-step guide")

        state[user_id].update({"step": "awaiting_approval", "doc_type": doc_type})
        say(f"On it — generating a *{doc_type}* now...")

        try:
            doc_data = generate_doc_json(s["content"], doc_type, s.get("source","slack"))
            state[user_id]["doc_data"] = doc_data
            send_doc_to_user(user_id, doc_data, say)
        except Exception as e:
            say(f"Something went wrong generating the doc: {str(e)}")
            state.pop(user_id, None)
        return

    # ── ON-DEMAND REQUEST (no existing state) ──
    # Detect if user is making a fresh request
    has_attachment  = bool(message.get("files"))
    mentions_create = any(w in text_lo for w in ["create", "write", "make", "generate", "build"])
    mentions_update = "update" in text_lo
    has_url         = bool(extract_url(text))

    if has_attachment or mentions_create or mentions_update or has_url or len(text) > 200:
        # Extract content
        content = text
        if has_attachment:
            # Try to get file content from attachment
            for f in message.get("files", []):
                if f.get("mimetype", "").startswith("text") or f.get("filetype") in ("txt","vtt","srt"):
                    try:
                        r = requests.get(f["url_private"], headers={"Authorization": f"Bearer {os.environ.get('SLACK_BOT_TOKEN')}"})
                        content += "\n\n" + r.text
                    except Exception:
                        pass

        # Detect doc type and action from message
        doc_type   = parse_doc_type(text) or "step-by-step guide"
        action     = "update" if mentions_update else "new"
        update_url = extract_url(text) if mentions_update else None
        source     = "zoom" if "zoom" in text_lo else "gong" if "gong" in text_lo else "slack"

        state[user_id] = {
            "step": "awaiting_approval",
            "content": content,
            "source": source,
            "action": action,
            "doc_type": doc_type,
            "update_url": update_url,
        }

        say(f"On it — generating a *{doc_type}* now...")
        try:
            doc_data = generate_doc_json(content, doc_type, source)
            state[user_id]["doc_data"] = doc_data
            send_doc_to_user(user_id, doc_data, say)
        except Exception as e:
            say(f"Something went wrong: {str(e)}")
            state.pop(user_id, None)
    else:
        say(
            "Hey! Here's what I can do:\n\n"
            "• *Paste any transcript or thread* and I'll turn it into a doc\n"
            "• *Create a [FAQ / step-by-step guide / release notes] from this* + paste content\n"
            "• *Update [article link] based on this* + paste content\n\n"
            "I'm also watching all your Slack channels and will ping you if I spot something worth documenting."
        )


# ── Zoom webhook ──────────────────────────────────────────────────────────────

@flask_app.route("/zoom/webhook", methods=["POST"])
def zoom_webhook():
    """
    Receives Zoom recording.completed events.
    Finds the host's Slack account by email and DMs them.
    """
    data = request.get_json(force=True)

    # Zoom sends a validation challenge on first setup
    if data.get("event") == "endpoint.url_validation":
        return jsonify({"plainToken": data["payload"]["plainToken"],
                        "encryptedToken": data["payload"]["plainToken"]})

    if data.get("event") != "recording.completed":
        return jsonify({"status": "ignored"}), 200

    obj         = data["payload"]["object"]
    topic       = obj.get("topic", "Untitled Meeting")
    host_email  = obj.get("host_email", "")
    files       = obj.get("recording_files", [])

    # Find transcript file
    transcript_text = ""
    for f in files:
        if f.get("file_type") == "TRANSCRIPT":
            try:
                r = requests.get(
                    f["download_url"],
                    headers={"Authorization": f"Bearer {os.environ.get('ZOOM_API_TOKEN','')}"},
                    timeout=30,
                )
                transcript_text = r.text
            except Exception:
                pass
            break

    if not transcript_text:
        return jsonify({"status": "no transcript"}), 200

    # Find host's Slack user by email
    try:
        result  = slack_app.client.users_lookupByEmail(email=host_email)
        user_id = result["user"]["id"]
    except Exception:
        return jsonify({"status": "host not found in Slack"}), 200

    # Analyse the transcript
    analysis = is_doc_worthy(transcript_text[:3000])

    state[user_id] = {
        "step":           "awaiting_action",
        "content":        transcript_text,
        "source":         "zoom",
        "suggested_type": analysis.get("suggested_type", "step-by-step guide"),
    }

    slack_app.client.chat_postMessage(
        channel=user_id,
        text=(
            f"Hey — VicSherlock here. I just analysed your Zoom recording: *{topic}*\n\n"
            f"{'*What I found:* ' + analysis.get('summary','') + chr(10) + chr(10) if analysis.get('summary') else ''}"
            f"Would you like me to create a doc from this?\n"
            f"• Reply *new* to create a new guide\n"
            f"• Reply *update* + paste an existing article link to update it\n"
            f"• Reply *skip* to ignore\n\n"
            f"I'd suggest a *{analysis.get('suggested_type','guide')}*."
        ),
    )

    return jsonify({"status": "ok"}), 200


# ── Slack event endpoint ──────────────────────────────────────────────────────

@flask_app.route("/slack/events", methods=["POST"])
def slack_events():
    return handler.handle(request)


# ── Health check ──────────────────────────────────────────────────────────────

@flask_app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "VicSherlock is watching"}), 200


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    port  = int(os.environ.get("PORT", 3000))
    flask_app.run(host="0.0.0.0", port=port, debug=False)
